import math
import random
import sys
import os
from datetime import datetime
from typing import List, Optional, Tuple

import numpy as np
import pandas as pd
import pygame

# Try to use gfxdraw; fall back gracefully if unavailable
try:
    import pygame.gfxdraw as gfx
    HAS_GFX = True
except Exception:
    gfx = None
    HAS_GFX = False

# Try to import python-docx for Word output
try:
    from docx import Document
    HAS_DOCX = True
except Exception:
    Document = None
    HAS_DOCX = False

# Optional JIT (auto if installed)
try:
    from numba import njit
    HAS_NUMBA = True
except Exception:
    HAS_NUMBA = False
    njit = None

from concurrent.futures import ThreadPoolExecutor, as_completed

# --------------------------
# Config & Constants
# --------------------------
BLUE   = (0, 0, 255)
BLACK  = (0, 0, 0)
RED    = (255, 0, 0)
YEL    = (255, 255, 0)
WHITE  = (255, 255, 255)
GREY   = (120, 120, 120)

ROW_COUNT     = 6
COLUMN_COUNT  = 7
WINDOW_LENGTH = 4

PLAYER = 0      # red
AI     = 1      # yellow (can be human in PvP)

EMPTY        = 0
PLAYER_PIECE = 1
AI_PIECE     = 2

# Depth slider (drag in the top bar)
SLIDER_MIN = 1
SLIDER_MAX = 9
SLIDER_W   = 220

WIN_SCORE  = 10_000_000
LOSE_SCORE = -10_000_000
DRAW_SCORE = 0

# Parallelization
USE_PARALLEL_ROOT = True
MAX_WORKERS = max(2, (os.cpu_count() or 4) // 2)  # gentle default

# UI
SQUARESIZE = 100
RADIUS     = SQUARESIZE // 2 - 6
WIDTH      = COLUMN_COUNT * SQUARESIZE
TOPBAR_H   = SQUARESIZE
HEIGHT     = TOPBAR_H + ROW_COUNT * SQUARESIZE
SIZE       = (WIDTH, HEIGHT)

# --------------------------
# Board helpers
# --------------------------
def create_board() -> np.ndarray:
    return np.zeros((ROW_COUNT, COLUMN_COUNT), dtype=np.int8)

def drop_piece(board: np.ndarray, row: int, col: int, piece: int) -> None:
    board[row, col] = piece

def is_valid_location(board: np.ndarray, col: int) -> bool:
    return board[ROW_COUNT - 1, col] == EMPTY

def get_next_open_row(board: np.ndarray, col: int) -> Optional[int]:
    for r in range(ROW_COUNT):
        if board[r, col] == EMPTY:
            return r
    return None

def get_valid_locations(board: np.ndarray) -> List[int]:
    return [c for c in range(COLUMN_COUNT) if is_valid_location(board, c)]

def print_board(board: np.ndarray) -> None:
    print(np.flip(board, 0))

# --------------------------
# Game state checks (Python)
# --------------------------
def winning_move_py(board: np.ndarray, piece: int) -> bool:
    # Horizontal
    for c in range(COLUMN_COUNT - 3):
        for r in range(ROW_COUNT):
            if (board[r, c] == piece and board[r, c+1] == piece
                and board[r, c+2] == piece and board[r, c+3] == piece):
                return True
    # Vertical
    for c in range(COLUMN_COUNT):
        for r in range(ROW_COUNT - 3):
            if (board[r, c] == piece and board[r+1, c] == piece
                and board[r+2, c] == piece and board[r+3, c] == piece):
                return True
    # Diagonal \
    for c in range(COLUMN_COUNT - 3):
        for r in range(ROW_COUNT - 3):
            if (board[r, c] == piece and board[r+1, c+1] == piece
                and board[r+2, c+2] == piece and board[r+3, c+3] == piece):
                return True
    # Diagonal /
    for c in range(COLUMN_COUNT - 3):
        for r in range(3, ROW_COUNT):
            if (board[r, c] == piece and board[r-1, c+1] == piece
                and board[r-2, c+2] == piece and board[r-3, c+3] == piece):
                return True
    return False

def evaluate_window_py(window: List[int], piece: int) -> int:
    score = 0
    opp = PLAYER_PIECE if piece == AI_PIECE else AI_PIECE
    if window.count(piece) == 4: score += 100
    elif window.count(piece) == 3 and window.count(EMPTY) == 1: score += 5
    elif window.count(piece) == 2 and window.count(EMPTY) == 2: score += 2
    if window.count(opp) == 3 and window.count(EMPTY) == 1: score -= 4
    return score

def score_position_py(board: np.ndarray, piece: int) -> int:
    score = 0
    center_array = [int(i) for i in list(board[:, COLUMN_COUNT // 2])]
    score += center_array.count(piece) * 3
    for r in range(ROW_COUNT):
        row_array = [int(i) for i in list(board[r, :])]
        for c in range(COLUMN_COUNT - 3):
            score += evaluate_window_py(row_array[c:c+WINDOW_LENGTH], piece)
    for c in range(COLUMN_COUNT):
        col_array = [int(i) for i in list(board[:, c])]
        for r in range(ROW_COUNT - 3):
            score += evaluate_window_py(col_array[r:r+WINDOW_LENGTH], piece)
    for r in range(ROW_COUNT - 3):
        for c in range(COLUMN_COUNT - 3):
            score += evaluate_window_py([int(board[r+i, c+i]) for i in range(WINDOW_LENGTH)], piece)
    for r in range(ROW_COUNT - 3):
        for c in range(COLUMN_COUNT - 3):
            score += evaluate_window_py([int(board[r+3-i, c+i]) for i in range(WINDOW_LENGTH)], piece)
    return score

# --------------------------
# Optional JIT versions
# --------------------------
if HAS_NUMBA:
    @njit(cache=True)
    def winning_move_njit(board: np.ndarray, piece: int) -> bool:
        rows, cols = board.shape
        # Horizontal
        for c in range(cols - 3):
            for r in range(rows):
                if (board[r, c] == piece and board[r, c+1] == piece
                    and board[r, c+2] == piece and board[r, c+3] == piece):
                    return True
        # Vertical
        for c in range(cols):
            for r in range(rows - 3):
                if (board[r, c] == piece and board[r+1, c] == piece
                    and board[r+2, c] == piece and board[r+3, c] == piece):
                    return True
        # Diagonal \
        for c in range(cols - 3):
            for r in range(rows - 3):
                if (board[r, c] == piece and board[r+1, c+1] == piece
                    and board[r+2, c+2] == piece and board[r+3, c+3] == piece):
                    return True
        # Diagonal /
        for c in range(cols - 3):
            for r in range(3, rows):
                if (board[r, c] == piece and board[r-1, c+1] == piece
                    and board[r-2, c+2] == piece and board[r-3, c+3] == piece):
                    return True
        return False

    @njit(cache=True)
    def score_position_njit(board: np.ndarray, piece: int) -> int:
        rows, cols = board.shape
        score = 0
        center_col = cols // 2
        # center
        cnt = 0
        for i in range(rows):
            if board[i, center_col] == piece:
                cnt += 1
        score += 3 * cnt

        # rows
        for r in range(rows):
            for c in range(cols - 3):
                w0 = board[r, c]
                w1 = board[r, c+1]
                w2 = board[r, c+2]
                w3 = board[r, c+3]
                score += _eval_window_njit(w0, w1, w2, w3, piece)
        # cols
        for c in range(cols):
            for r in range(rows - 3):
                w0 = board[r, c]
                w1 = board[r+1, c]
                w2 = board[r+2, c]
                w3 = board[r+3, c]
                score += _eval_window_njit(w0, w1, w2, w3, piece)
        # diag \
        for r in range(rows - 3):
            for c in range(cols - 3):
                w0 = board[r, c]
                w1 = board[r+1, c+1]
                w2 = board[r+2, c+2]
                w3 = board[r+3, c+3]
                score += _eval_window_njit(w0, w1, w2, w3, piece)
        # diag /
        for r in range(rows - 3):
            for c in range(cols - 3):
                w0 = board[r+3, c]
                w1 = board[r+2, c+1]
                w2 = board[r+1, c+2]
                w3 = board[r,   c+3]
                score += _eval_window_njit(w0, w1, w2, w3, piece)
        return score

    @njit(cache=True)
    def _eval_window_njit(a, b, c, d, piece) -> int:
        score = 0
        opp = 1 if piece == 2 else 2
        cnt_p = int(a == piece) + int(b == piece) + int(c == piece) + int(d == piece)
        cnt_e = int(a == 0) + int(b == 0) + int(c == 0) + int(d == 0)
        cnt_o = int(a == opp)  + int(b == opp)  + int(c == opp)  + int(d == opp)
        if cnt_p == 4: score += 100
        elif cnt_p == 3 and cnt_e == 1: score += 5
        elif cnt_p == 2 and cnt_e == 2: score += 2
        if cnt_o == 3 and cnt_e == 1: score -= 4
        return score

# Bind function pointers (JIT if available)
WIN_CHECK   = winning_move_njit   if HAS_NUMBA else winning_move_py
SCORE_FUNC  = score_position_njit if HAS_NUMBA else score_position_py

def winning_move(board: np.ndarray, piece: int) -> bool:
    # wrapper to keep bool type consistent across njit/py
    return bool(WIN_CHECK(board, piece))

def score_position(board: np.ndarray, piece: int) -> int:
    return int(SCORE_FUNC(board, piece))

def is_terminal_node(board: np.ndarray) -> bool:
    return winning_move(board, PLAYER_PIECE) or winning_move(board, AI_PIECE) or len(get_valid_locations(board)) == 0

# --------------------------
# Minimax (with ordering)
# --------------------------
def ordered_valid_locations(board: np.ndarray) -> List[int]:
    # prefer center
    order = [3,2,4,1,5,0,6]
    return [c for c in order if is_valid_location(board, c)]

def minimax(board: np.ndarray, depth: int, alpha: float, beta: float, maximizing: bool) -> Tuple[Optional[int], int]:
    valid_locations = ordered_valid_locations(board)
    terminal = is_terminal_node(board)
    if depth == 0 or terminal:
        if terminal:
            if winning_move(board, AI_PIECE):       return None, WIN_SCORE
            elif winning_move(board, PLAYER_PIECE): return None, LOSE_SCORE
            else:                                   return None, DRAW_SCORE
        return None, score_position(board, AI_PIECE)

    best_col = valid_locations[0] if valid_locations else None
    if maximizing:
        value = -math.inf
        for col in valid_locations:
            row = get_next_open_row(board, col)
            if row is None: continue
            b_copy = board.copy()
            drop_piece(b_copy, row, col, AI_PIECE)
            _, new_score = minimax(b_copy, depth-1, alpha, beta, False)
            if new_score > value:
                value, best_col = new_score, col
            alpha = max(alpha, value)
            if alpha >= beta: break
        return best_col, int(value)
    else:
        value = math.inf
        for col in valid_locations:
            row = get_next_open_row(board, col)
            if row is None: continue
            b_copy = board.copy()
            drop_piece(b_copy, row, col, PLAYER_PIECE)
            _, new_score = minimax(b_copy, depth-1, alpha, beta, True)
            if new_score < value:
                value, best_col = new_score, col
            beta = min(beta, value)
            if alpha >= beta: break
        return best_col, int(value)

# ---------- Root parallelization ----------
def _minimax_child(board: np.ndarray, col: int, depth: int) -> Tuple[int, int]:
    """Evaluate one child move for AI root (maximizing)."""
    row = get_next_open_row(board, col)
    if row is None:
        return col, -10**9
    b_copy = board.copy()
    drop_piece(b_copy, row, col, AI_PIECE)
    _, sc = minimax(b_copy, depth-1, -math.inf, math.inf, False)
    return col, sc

def minimax_root_parallel(board: np.ndarray, depth: int, maximizing: bool, candidate_moves: List[int]) -> Tuple[Optional[int], int]:
    if not candidate_moves:
        return None, 0
    if not USE_PARALLEL_ROOT or len(candidate_moves) == 1 or depth <= 2:
        # sequential fallback
        best_col, best_score = None, -math.inf
        for c in candidate_moves:
            col, sc = _minimax_child(board, c, depth)
            # center tie-break
            key = (sc, -abs(c - COLUMN_COUNT//2))
            if best_col is None or key > (best_score, -abs(best_col - COLUMN_COUNT//2)):
                best_col, best_score = col, sc
        return best_col, int(best_score)

    best_col, best_score = None, -math.inf
    with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, len(candidate_moves))) as ex:
        futures = {ex.submit(_minimax_child, board, c, depth): c for c in candidate_moves}
        for fut in as_completed(futures):
            col, sc = fut.result()
            key = (sc, -abs(col - COLUMN_COUNT//2))
            if best_col is None or key > (best_score, -abs(best_col - COLUMN_COUNT//2)):
                best_col, best_score = col, sc
    return best_col, int(best_score)

# --------------------------
# Tactical helpers to make AI smarter
# --------------------------
def immediate_wins(board: np.ndarray, piece: int) -> List[int]:
    cols = []
    for c in get_valid_locations(board):
        r = get_next_open_row(board, c)
        if r is None: 
            continue
        b2 = board.copy()
        drop_piece(b2, r, c, piece)
        if winning_move(b2, piece):
            cols.append(c)
    return cols

def is_unsafe_for_ai(board: np.ndarray, ai_col: int) -> bool:
    """After AI plays ai_col, can the human win immediately next turn?"""
    r = get_next_open_row(board, ai_col)
    if r is None:
        return True
    b2 = board.copy()
    drop_piece(b2, r, ai_col, AI_PIECE)
    # human immediate wins now?
    return len(immediate_wins(b2, PLAYER_PIECE)) > 0

# --------------------------
# Drawing helpers (centers, slider, glow, controls)
# --------------------------
def rc_to_screen(col: int, row: int) -> Tuple[int, int]:
    # board row 0 is bottom -> flip vertically for drawing
    x = int(col * SQUARESIZE + SQUARESIZE / 2)
    y = HEIGHT - int(row * SQUARESIZE + SQUARESIZE / 2)
    return x, y

def aa_ring(surface: pygame.Surface, center: Tuple[int,int], color: Tuple[int,int,int], radius: int, width: int=3):
    if HAS_GFX:
        for w in range(width):
            gfx.aacircle(surface, center[0], center[1], radius + w, color)
            gfx.circle(surface,  center[0], center[1], radius + w, color)
    else:
        pygame.draw.circle(surface, color, center, radius + width//2, width)

def glow_disc(surface: pygame.Surface, center: Tuple[int,int], base_color: Tuple[int,int,int]):
    glow = pygame.Surface((SQUARESIZE*2, SQUARESIZE*2), pygame.SRCALPHA)
    gx, gy = glow.get_width()//2, glow.get_height()//2
    for i in range(6, 0, -1):
        r = RADIUS + 3 + i*2
        alpha = 22 + i*6
        pygame.draw.circle(glow, (*base_color, alpha), (gx, gy), r)
    surface.blit(glow, glow.get_rect(center=center))
    aa_ring(surface, center, base_color, RADIUS + 4, width=3)

# Slider geometry
def slider_rects(ai_depth: int) -> Tuple[pygame.Rect, Tuple[int,int], int]:
    w = SLIDER_W
    x = WIDTH - w - 20
    y = TOPBAR_H - 28
    track = pygame.Rect(x, y, w, 8)
    handle_r = 10
    ratio = (ai_depth - SLIDER_MIN) / (SLIDER_MAX - SLIDER_MIN)
    ratio = max(0.0, min(1.0, ratio))
    hx = x + int(ratio * w)
    hy = y + track.height // 2
    return track, (hx, hy), handle_r

def x_to_depth(x: int, track: pygame.Rect) -> int:
    ratio = (x - track.left) / max(1, track.width)
    val = SLIDER_MIN + round(max(0.0, min(1.0, ratio)) * (SLIDER_MAX - SLIDER_MIN))
    return max(SLIDER_MIN, min(SLIDER_MAX, val))

def draw_slider(surface: pygame.Surface, ai_depth: int, active: bool = True):
    track, (hx, hy), hr = slider_rects(ai_depth)
    # track + fill
    pygame.draw.rect(surface, (60,60,60) if active else (40,40,40), track, border_radius=4)
    fill = pygame.Rect(track.left, track.top, max(0, hx - track.left), track.height)
    pygame.draw.rect(surface, (200,200,200) if active else (90,90,90), fill, border_radius=4)
    # handle
    handle_col = WHITE if active else GREY
    if HAS_GFX:
        gfx.filled_circle(surface, hx, hy, hr, handle_col)
        gfx.aacircle(surface, hx, hy, hr, (220,220,220))
    else:
        pygame.draw.circle(surface, handle_col, (hx, hy), hr)
    # label
    font = pygame.font.SysFont("monospace", 18, bold=True)
    label_col = WHITE if active else GREY
    surface.blit(font.render(f"Depth: {ai_depth}", True, label_col), (track.left - 110, track.top - 4))

# Control buttons (mode + first)
def draw_button(surface: pygame.Surface, rect: pygame.Rect, text: str, enabled: bool=True):
    col = WHITE if enabled else GREY
    pygame.draw.rect(surface, (25,25,25), rect, border_radius=8)
    pygame.draw.rect(surface, col, rect, width=2, border_radius=8)
    font = pygame.font.SysFont("monospace", 20, bold=True)
    tw, th = font.size(text)
    surface.blit(font.render(text, True, col), (rect.x + (rect.w - tw)//2, rect.y + (rect.h - th)//2))

def get_control_rects():
    mode_rect  = pygame.Rect(12, TOPBAR_H - 36, 180, 28)
    first_rect = pygame.Rect(200, TOPBAR_H - 36, 160, 28)
    return mode_rect, first_rect

# --------------------------
# UI drawing
# --------------------------
def draw_board(surface: pygame.Surface,
               board: np.ndarray,
               last_user: Optional[Tuple[int,int]],
               last_ai: Optional[Tuple[int,int]],
               msg: str,
               ai_depth: int,
               mode: str,
               first: str) -> None:
    # grid
    for c in range(COLUMN_COUNT):
        for r in range(ROW_COUNT):
            pygame.draw.rect(surface, BLUE, (c*SQUARESIZE, r*SQUARESIZE+TOPBAR_H, SQUARESIZE, SQUARESIZE))
            pygame.draw.circle(surface, BLACK,
                (int(c*SQUARESIZE + SQUARESIZE/2), int(r*SQUARESIZE + TOPBAR_H + SQUARESIZE/2)),
                RADIUS)

    # pieces
    for c in range(COLUMN_COUNT):
        for r in range(ROW_COUNT):
            if board[r, c] == PLAYER_PIECE:
                pygame.draw.circle(surface, RED, rc_to_screen(c, r), RADIUS)
            elif board[r, c] == AI_PIECE:
                pygame.draw.circle(surface, YEL, rc_to_screen(c, r), RADIUS)

    # glow for last moves
    if last_user:
        glow_disc(surface, rc_to_screen(last_user[1], last_user[0]), RED)
    if last_ai:
        glow_disc(surface, rc_to_screen(last_ai[1], last_ai[0]), YEL)

    # top bar
    pygame.draw.rect(surface, BLACK, (0, 0, WIDTH, TOPBAR_H))
    font = pygame.font.SysFont("monospace", 32, bold=True)
    surface.blit(font.render(msg, True, WHITE if "Your" in msg or "New" in msg else YEL), (12, 14))

    # controls
    mode_rect, first_rect = get_control_rects()
    draw_button(surface, mode_rect, f"Mode: {mode}", True)
    first_enabled = (mode == "PvE")
    draw_button(surface, first_rect, f"First: {first}", first_enabled)

    # slider
    draw_slider(surface, ai_depth, active=(mode == "PvE"))

    pygame.display.update()

def show_message(surface: pygame.Surface, text: str, ai_depth: int, mode: str, first: str) -> None:
    pygame.draw.rect(surface, BLACK, (0, 0, WIDTH, TOPBAR_H))
    font = pygame.font.SysFont("monospace", 48, bold=True)
    label = font.render(text, True, YEL if "AI" in text else WHITE)
    rect = label.get_rect(midleft=(12, TOPBAR_H//2))
    surface.blit(label, rect)
    # controls
    mode_rect, first_rect = get_control_rects()
    draw_button(surface, mode_rect, f"Mode: {mode}", True)
    draw_button(surface, first_rect, f"First: {first}", mode == "PvE")
    draw_slider(surface, ai_depth, active=(mode == "PvE"))
    pygame.display.update()

# --------------------------
# Drop animation
# --------------------------
def draw_filled_disc(surface, x, y, color):
    if HAS_GFX:
        gfx.filled_circle(surface, x, y, RADIUS, color)
        gfx.aacircle(surface, x, y, RADIUS, color)
    else:
        pygame.draw.circle(surface, color, (x, y), RADIUS)

def animate_drop(screen: pygame.Surface,
                 board: np.ndarray,
                 col: int,
                 row: int,
                 color: Tuple[int,int,int],
                 last_user: Optional[Tuple[int,int]],
                 last_ai: Optional[Tuple[int,int]],
                 msg: str,
                 ai_depth: int,
                 mode: str,
                 first: str):
    clock = pygame.time.Clock()
    sx, sy = col * SQUARESIZE + SQUARESIZE//2, TOPBAR_H//2
    ex, ey = rc_to_screen(col, row)
    y = sy
    v = 0.0
    g = 1.25
    while y < ey - 1:
        for e in pygame.event.get():
            if e.type == pygame.QUIT:
                pygame.quit(); sys.exit()
        v = min(v + g, 28)
        y = min(y + v, ey)
        draw_board(screen, board, last_user, last_ai, msg, ai_depth, mode, first)
        draw_filled_disc(screen, sx, int(y), color)
        pygame.display.update()
        clock.tick(120)

# --------------------------
# Word logging (Pandas -> DOCX)
# --------------------------
def save_game_result(board: np.ndarray, result: str, docx_path: str = "Connect4_Game_Logs.docx") -> None:
    """
    Append the final board to a Word document using Pandas.
    - Adds a table with symbols (R=Player, Y=AI, ·=empty), top row first.
    - Also adds a list-of-lists numeric matrix for reference.
    Falls back to CSV if python-docx is not installed.
    """
    df_num = pd.DataFrame(
        np.flip(board, 0).astype(int),
        columns=[f"C{c+1}" for c in range(COLUMN_COUNT)]
    )
    df_sym = df_num.replace({0: "·", 1: "R", 2: "Y"})
    stamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")

    if not HAS_DOCX:
        csv_path = "Connect4_Game_Logs.csv"
        with open(csv_path, "a", encoding="utf-8") as f:
            f.write(f"# {stamp} — {result}\n")
        df_num.to_csv(csv_path, mode="a", index=False)
        with open(csv_path, "a", encoding="utf-8") as f:
            f.write("\n")
        print("[INFO] python-docx not found; wrote to Connect4_Game_Logs.csv instead.")
        return

    doc = Document(docx_path) if os.path.exists(docx_path) else Document()
    doc.add_heading(f"Game Result — {result}", level=1)
    doc.add_paragraph(stamp)
    doc.add_paragraph("Legend: R = Player (Red), Y = AI/Player2 (Yellow), · = empty")

    rows, cols = df_sym.shape
    table = doc.add_table(rows=rows + 1, cols=cols + 1)
    hdr = table.rows[0].cells
    hdr[0].text = ""
    for j, cname in enumerate(df_sym.columns, start=1):
        hdr[j].text = cname
    for i in range(rows):
        cells = table.rows[i + 1].cells
        cells[0].text = f"R{ROW_COUNT - i}"
        for j in range(cols):
            cells[j + 1].text = str(df_sym.iat[i, j])

    doc.add_paragraph("Numeric matrix (top row first):")
    doc.add_paragraph(str(df_num.values.tolist()))
    doc.add_paragraph("")  # spacer

    doc.save(docx_path)
    print(f"[INFO] Saved board to {docx_path}")

# --------------------------
# Best-move chooser (no artificial delay)
# --------------------------
def choose_best_ai_move(board: np.ndarray, depth: int) -> int:
    # 1) Win immediately if possible
    wins = immediate_wins(board, AI_PIECE)
    if wins:
        # prefer closer to center among winning moves
        return max(wins, key=lambda c: -abs(c - COLUMN_COUNT//2))
    # 2) Block opponent's immediate win
    opp_wins = immediate_wins(board, PLAYER_PIECE)
    if opp_wins:
        return max(opp_wins, key=lambda c: -abs(c - COLUMN_COUNT//2))
    # 3) Avoid obvious blunders (opponent wins next)
    valids = get_valid_locations(board)
    safe = [c for c in valids if not is_unsafe_for_ai(board, c)]
    candidate_moves = safe if safe else valids
    # 4) Parallelized root minimax over candidates
    col, _ = minimax_root_parallel(board, depth, True, candidate_moves)
    return col if col is not None else candidate_moves[0]

# --------------------------
# Game loop
# --------------------------
def main() -> None:
    pygame.init()
    pygame.display.set_caption("Connect Four (Player vs AI / PvP)")
    screen = pygame.display.set_mode(SIZE)
    clock  = pygame.time.Clock()

    board = create_board()
    last_user: Optional[Tuple[int,int]] = None  # (row, col) last red
    last_bot:  Optional[Tuple[int,int]] = None  # (row, col) last yellow

    mode = "PvE"      # "PvE" (Human vs AI)  or "PvP" (Human vs Human)
    first = "Human"   # In PvE: "Human" or "AI". In PvP: Red starts.
    ai_depth = 5      # default
    slider_drag = False

    # who starts?
    turn = (AI if (mode == "PvE" and first == "AI") else PLAYER)
    game_over = False

    # initialize hover to actual mouse x so first click is correct
    hover_x = pygame.mouse.get_pos()[0]

    draw_board(screen, board, last_user, last_bot,
               "Your turn — click a column" if turn == PLAYER else ("AI thinking…" if mode == "PvE" else "Yellow's turn"),
               ai_depth, mode, first)

    while True:
        clock.tick(60)

        for event in pygame.event.get():
            if event.type == pygame.QUIT:
                pygame.quit(); sys.exit()

            if event.type == pygame.KEYDOWN:
                if event.key in (pygame.K_ESCAPE, pygame.K_q):
                    pygame.quit(); sys.exit()
                if event.key == pygame.K_r:
                    board[:] = 0
                    last_user = last_bot = None
                    game_over = False
                    turn = (AI if (mode == "PvE" and first == "AI") else PLAYER)
                    draw_board(screen, board, last_user, last_bot,
                               "New game! Your turn" if turn == PLAYER else ("AI thinking…" if mode == "PvE" else "Yellow's turn"),
                               ai_depth, mode, first)

            # Top bar interactions (slider + buttons)
            if event.type == pygame.MOUSEBUTTONDOWN:
                mx, my = event.pos
                if my < TOPBAR_H:
                    # Buttons
                    mode_rect, first_rect = get_control_rects()
                    if mode_rect.collidepoint(mx, my):
                        # Toggle mode
                        mode = "PvP" if mode == "PvE" else "PvE"
                        # Reset game with new mode
                        board[:] = 0
                        last_user = last_bot = None
                        game_over = False
                        if mode == "PvE":
                            # keep 'first' as-is; decide who starts
                            turn = (AI if first == "AI" else PLAYER)
                        else:
                            # PvP: red always first
                            turn = PLAYER
                        draw_board(screen, board, last_user, last_bot,
                                   "New game! Your turn" if turn == PLAYER else ("AI thinking…" if mode == "PvE" else "Yellow's turn"),
                                   ai_depth, mode, first)
                        continue
                    if first_rect.collidepoint(mx, my) and mode == "PvE":
                        first = "AI" if first == "Human" else "Human"
                        # Restart to apply who starts
                        board[:] = 0
                        last_user = last_bot = None
                        game_over = False
                        turn = (AI if first == "AI" else PLAYER)
                        draw_board(screen, board, last_user, last_bot,
                                   "New game! Your turn" if turn == PLAYER else "AI thinking…",
                                   ai_depth, mode, first)
                        continue

                    # Slider
                    track, (hx, hy), hr = slider_rects(ai_depth)
                    if mode == "PvE" and (track.inflate(20, 16).collidepoint(mx, my) or (mx - hx)**2 + (my - hy)**2 <= (hr+6)**2):
                        slider_drag = True
                        ai_depth = x_to_depth(mx, track)
                        draw_board(screen, board, last_user, last_bot, "Depth changed", ai_depth, mode, first)
                        continue  # don't treat as board click

            if event.type == pygame.MOUSEBUTTONUP:
                slider_drag = False

            if event.type == pygame.MOUSEMOTION and slider_drag and mode == "PvE":
                mx, my = event.pos
                track, _, _ = slider_rects(ai_depth)
                ai_depth = x_to_depth(mx, track)
                draw_board(screen, board, last_user, last_bot, "Depth changed", ai_depth, mode, first)
                continue

            if game_over:
                continue

            # Hover preview in top bar
            if event.type == pygame.MOUSEMOTION:
                hover_x = max(0, min(event.pos[0], WIDTH))
                if event.pos[1] < TOPBAR_H:
                    pygame.draw.rect(screen, BLACK, (0, 0, WIDTH, TOPBAR_H))
                    pygame.draw.circle(screen, RED if turn == PLAYER else YEL, (hover_x, TOPBAR_H // 2), RADIUS)
                    # controls & slider redraw
                    mode_rect, first_rect = get_control_rects()
                    draw_button(screen, mode_rect, f"Mode: {mode}", True)
                    draw_button(screen, first_rect, f"First: {first}", mode == "PvE")
                    draw_slider(screen, ai_depth, active=(mode == "PvE"))
                    pygame.display.update()

            # Human click on board
            if event.type == pygame.MOUSEBUTTONDOWN:
                mx, my = event.pos
                if my < TOPBAR_H:
                    continue  # click on top bar

                # choose column from the actual click X (not hover_x)
                col = mx // SQUARESIZE

                if not is_valid_location(board, col):
                    show_message(screen, "Column full. Try another.", ai_depth, mode, first)
                    continue

                row = get_next_open_row(board, col)
                if row is None:
                    continue

                # Decide piece/color by whose turn (PvP or PvE-human)
                piece_color = RED if turn == PLAYER else YEL
                piece_code  = PLAYER_PIECE if turn == PLAYER else AI_PIECE

                animate_drop(screen, board, col, row, piece_color, last_user, last_bot,
                             "Your move…" if turn == PLAYER else ("Yellow move…" if mode=="PvP" else "AI thinking…"),
                             ai_depth, mode, first)
                drop_piece(board, row, col, piece_code)

                if turn == PLAYER:
                    last_user = (row, col)
                else:
                    last_bot = (row, col)

                print_board(board)
                draw_board(screen, board, last_user, last_bot, "…" , ai_depth, mode, first)

                # Win/Draw check
                if winning_move(board, piece_code):
                    if mode == "PvE":
                        if turn == PLAYER:
                            show_message(screen, "You win! 🎉", ai_depth, mode, first)
                            save_game_result(board, "Player (RED) wins vs AI")
                        else:
                            show_message(screen, "AI wins!", ai_depth, mode, first)
                            save_game_result(board, "AI (YELLOW) wins vs Player")
                    else:
                        # PvP: announce by color
                        msg = "Red wins! 🎉" if turn == PLAYER else "Yellow wins!"
                        show_message(screen, msg, ai_depth, mode, first)
                        save_game_result(board, msg)
                    game_over = True
                    continue

                if len(get_valid_locations(board)) == 0:
                    show_message(screen, "Draw!", ai_depth, mode, first)
                    save_game_result(board, "Draw")
                    game_over = True
                    continue

                # Next turn
                if mode == "PvE":
                    # human -> AI
                    if turn == PLAYER:
                        turn = AI
                        show_message(screen, "AI thinking…", ai_depth, mode, first)
                    else:
                        # shouldn't happen; AI handled below
                        turn = PLAYER
                        show_message(screen, "Your turn", ai_depth, mode, first)
                else:
                    # PvP: toggle between players
                    turn = PLAYER if turn == AI else AI
                    who = "Red's turn" if turn == PLAYER else "Yellow's turn"
                    show_message(screen, who, ai_depth, mode, first)

        # AI move (only in PvE) — NO artificial delays
        if not game_over and mode == "PvE" and turn == AI:
            col = choose_best_ai_move(board, ai_depth)

            if col is None:
                show_message(screen, "Draw!", ai_depth, mode, first)
                save_game_result(board, "Draw")
                game_over = True
            else:
                row = get_next_open_row(board, col)
                if row is not None:
                    animate_drop(screen, board, col, row, YEL, last_user, last_bot, "AI thinking…", ai_depth, mode, first)
                    drop_piece(board, row, col, AI_PIECE)
                    last_bot = (row, col)
                    print_board(board)
                    draw_board(screen, board, last_user, last_bot, "AI moved", ai_depth, mode, first)

                    if winning_move(board, AI_PIECE):
                        show_message(screen, "AI wins!", ai_depth, mode, first)
                        save_game_result(board, "AI (YELLOW) wins vs Player")
                        game_over = True
                    elif len(get_valid_locations(board)) == 0:
                        show_message(screen, "Draw!", ai_depth, mode, first)
                        save_game_result(board, "Draw")
                        game_over = True
                    else:
                        turn = PLAYER
                        show_message(screen, "Your turn", ai_depth, mode, first)

if __name__ == "__main__":
    main()
